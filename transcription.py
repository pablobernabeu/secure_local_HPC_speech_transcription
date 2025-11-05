#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Enhanced transcription script with:
1. Optional name masking
2. Comprehensive spelling corrections
3. Repetition pattern detection
4. Improved formatting
"""

import os
import sys
import re
import tempfile
import argparse
from pathlib import Path
import textwrap
from datetime import datetime
import warnings
import csv
import urllib.request
import json
# Make docx import optional for testing name masking
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    print("Warning: python-docx not available - Word document creation disabled")
    DOCX_AVAILABLE = False
    Document = None
warnings.filterwarnings("ignore")

# Audio processing imports - make conditional to allow name masking to work independently
try:
    import librosa
    import soundfile as sf
    import numpy as np
    import torch
    AUDIO_LIBS_AVAILABLE = True
except ImportError as e:
    print(f"Warning: Audio processing libraries not available - {e}")
    print("Name masking will work, but audio transcription will be disabled")
    AUDIO_LIBS_AVAILABLE = False
    librosa = None
    sf = None
    np = None
    torch = None

# Set FFmpeg path for audio loading (only if audio libs available)
if AUDIO_LIBS_AVAILABLE:
    try:
        import imageio_ffmpeg
        os.environ["IMAGEIO_FFMPEG_EXE"] = imageio_ffmpeg.get_ffmpeg_exe()
    except ImportError:
        print("Warning: imageio_ffmpeg not available")

# Transformers import - also conditional
try:
    from transformers import pipeline
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    print("Warning: transformers not available - transcription disabled")
    TRANSFORMERS_AVAILABLE = False
    pipeline = None

# Enhanced name masking with global database and logging
# Enhanced name masking with global database and logging
def ensure_names_dataset():
    """Ensure names-dataset package is available with enhanced path handling."""
    import sys
    import os
    
    # Add user site-packages to path if not already there
    try:
        import site
        user_site = site.getusersitepackages()
        if user_site not in sys.path:
            sys.path.insert(0, user_site)
            print(f"   📁 Added user site-packages to Python path: {user_site}")
    except:
        pass
    
    try:
        from names_dataset import NameDataset
        print("   ✅ names-dataset package found and imported")
        
        # Test if it actually works and discover API
        try:
            nd = NameDataset()
            
            # Test with the actual API - check first_names attribute
            test_passed = False
            
            if hasattr(nd, 'first_names'):
                try:
                    names_dict = nd.first_names
                    if isinstance(names_dict, dict) and len(names_dict) > 1000:
                        print(f"   ✅ Database test passed: {len(names_dict)} names available via first_names attribute")
                        test_passed = True
                except Exception as e:
                    print(f"   ⚠️  Error accessing first_names: {e}")
            
            # Fallback tests for other API methods
            if not test_passed:
                if hasattr(nd, 'search_first_name'):
                    try:
                        test_result = nd.search_first_name('john')
                        print(f"✅ Database test passed (search_first_name): john -> {test_result}")
                        test_passed = True
                    except:
                        pass
                elif hasattr(nd, 'search'):
                    try:
                        test_result = nd.search('john')
                        print(f"✅ Database test passed (search): john -> {test_result}")
                        test_passed = True
                    except:
                        pass
            
            if test_passed:
                print("   ✅ Global names database functional")
                return True
            else:
                available_methods = [m for m in dir(nd) if not m.startswith('_') and callable(getattr(nd, m))]
                print(f"   ⚠️  Could not find working API method. Available: {available_methods[:5]}...")
                return False
                
        except Exception as e:
            print(f"   ⚠️  Database initialization failed: {e}")
            print("   📋 Package available but database not accessible")
            return False
            
    except ImportError as e:
        print(f"   ⚠️  names-dataset package not available: {e}")
        print("   📋 Will use enhanced basic names fallback")
        return False
    except Exception as e:
        print(f"   ⚠️  Error loading names-dataset: {e}")
        print("   📋 Will use enhanced basic names fallback")
        return False

def load_global_names(use_facebook_names=False, use_facebook_surnames=False, selected_languages=None, exclude_common_words=True):
    """Load and filter global names database using names-dataset package.
    
    Args:
        use_facebook_names (bool): If True, use Facebook's global first names database (730K+ names).
                                 If False, use curated multilingual first names.
                                 Default: False (to avoid false positives)
        use_facebook_surnames (bool): If True, use Facebook's global surnames database (980K+ names).
                                    If False, use curated multilingual surnames.
                                    Default: False (to avoid false positives)
        selected_languages (list): Languages to include from curated database.
                                 If None, uses all available languages.
                                 Supported: ['english', 'chinese', 'french', 'german', 'hindi', 'spanish']
    """
    
    # Determine what to load based on flags
    use_any_facebook = use_facebook_names or use_facebook_surnames
    
    # If no Facebook database is requested, use multilingual curated names directly
    if not use_any_facebook:
        print("   [DATABASE] Using MULTILINGUAL CURATED NAMES DATABASE (Facebook database disabled)")
        print("   [INFO] Safer option with ~200 carefully selected names per language")
        print("   [SUCCESS] Reduces false positives from global names")
        return create_multilingual_curated_names(selected_languages, exclude_common_words)
    
    # Facebook database requested for names and/or surnames
    print("   [GLOBAL] FACEBOOK GLOBAL DATABASE REQUESTED")
    if use_facebook_names and use_facebook_surnames:
        print("   [LOAD] Loading both first names and surnames from Facebook database")
    elif use_facebook_names:
        print("   [LOAD] Loading first names only from Facebook database") 
    elif use_facebook_surnames:
        print("   [LOAD] Loading surnames only from Facebook database")
    print("   [WARNING] May cause false positives with common words")
    
    # Try to use the names-dataset package
    if ensure_names_dataset():
        try:
            print("   Loading names-dataset package...")
            from names_dataset import NameDataset
            
            # Initialize the dataset (this takes a moment)
            nd = NameDataset()
            print("   ✅ Names-dataset package loaded successfully")
            
            # Extract first names and surnames based on flags
            all_first_names = set()
            all_surnames = set()
            
            # Load first names if requested
            if use_facebook_names and hasattr(nd, 'first_names'):
                try:
                    print("   Loading first names dictionary...")
                    names_dict = nd.first_names
                    if isinstance(names_dict, dict):
                        # The keys are the actual names
                        all_first_names.update(name.lower().strip() for name in names_dict.keys() if name.strip())
                        print(f"   ✅ Extracted {len(all_first_names)} first names")
                    else:
                        print(f"   ⚠️  first_names is not a dictionary: {type(names_dict)}")
                except Exception as e:
                    print(f"   ❌ Error accessing first_names: {e}")
            elif use_facebook_names:
                print("   ⚠️  First names requested but first_names attribute not available")
            
            # Load surnames if requested
            if use_facebook_surnames and hasattr(nd, 'last_names'):
                try:
                    print("   Loading surnames dictionary...")
                    surnames_dict = nd.last_names
                    if isinstance(surnames_dict, dict):
                        # The keys are the actual surnames
                        all_surnames.update(name.lower().strip() for name in surnames_dict.keys() if name.strip())
                        print(f"   ✅ Extracted {len(all_surnames)} surnames")
                    else:
                        print(f"   ⚠️  last_names is not a dictionary: {type(surnames_dict)}")
                except Exception as e:
                    print(f"   ❌ Error accessing last_names: {e}")
            elif use_facebook_surnames:
                print("   ⚠️  Surnames requested but last_names attribute not available")
            
            # Fallback methods if first_names doesn't work and was requested
            if use_facebook_names and len(all_first_names) == 0:
                print("   Trying fallback API methods...")
                
                # Method 2: Try get_first_names() if available
                if hasattr(nd, 'get_first_names'):
                    try:
                        print("   Using get_first_names() method...")
                        names_data = nd.get_first_names()
                        if isinstance(names_data, (list, set)):
                            all_first_names.update(name.lower() for name in names_data)
                        elif isinstance(names_data, dict):
                            # Flatten dictionary structure
                            for key, value in names_data.items():
                                if isinstance(value, (list, set)):
                                    all_first_names.update(name.lower() for name in value)
                                elif isinstance(value, str):
                                    all_first_names.add(value.lower())
                    except Exception as e:
                        print(f"   get_first_names() failed: {e}")
                
                # Method 3: Try data attribute
                elif hasattr(nd, 'data'):
                    try:
                        print("   Using data attribute...")
                        data = nd.data
                        if isinstance(data, dict):
                            # Look for first names in various possible structures
                            for key, value in data.items():
                                if 'first' in key.lower() or 'given' in key.lower():
                                    if isinstance(value, (list, set)):
                                        all_first_names.update(name.lower() for name in value)
                    except Exception as e:
                        print(f"   data attribute failed: {e}")
            
            # Check if we successfully extracted the requested data
            names_loaded = (not use_facebook_names) or len(all_first_names) > 100
            surnames_loaded = (not use_facebook_surnames) or len(all_surnames) > 100
            
            if names_loaded and surnames_loaded:
                print(f"   ✅ FACEBOOK DATABASE LOADED SUCCESSFULLY")
                print(f"   📊 Raw extraction: {len(all_first_names)} first names, {len(all_surnames)} surnames")
                
                # Apply filtering to both sets (only if they were loaded)
                filtered_first_names = filter_problematic_names(all_first_names) if use_facebook_names else set()
                filtered_surnames = filter_problematic_names(all_surnames) if use_facebook_surnames else set()
                
                # Merge with curated multilingual names for the types not requested from Facebook
                if not use_facebook_names or not use_facebook_surnames:
                    basic_first_names, basic_surnames = create_multilingual_curated_names(selected_languages, exclude_common_words)
                    if not use_facebook_names:
                        filtered_first_names = basic_first_names
                    if not use_facebook_surnames:
                        filtered_surnames = basic_surnames
                
                print(f"   [INFO] Final counts: {len(filtered_first_names)} first names, {len(filtered_surnames)} surnames")
                if use_facebook_names:
                    print(f"   [REMOVED] {len(all_first_names) - len(filtered_first_names)} common first names")
                if use_facebook_surnames:
                    print(f"   [REMOVED] {len(all_surnames) - len(filtered_surnames)} common surnames")
                
                # Return both sets as a tuple
                return filtered_first_names, filtered_surnames
            else:
                missing = []
                if use_facebook_names and len(all_first_names) <= 100:
                    missing.append("first names")
                if use_facebook_surnames and len(all_surnames) <= 100:
                    missing.append("surnames")
                print(f"   [ERROR] Insufficient data for: {', '.join(missing)}")
                raise Exception(f"Insufficient data extracted for: {', '.join(missing)}")
                
        except Exception as e:
            print(f"   [ERROR] FACEBOOK DATABASE FAILED: {e}")
            print(f"   [FALLBACK] Using multilingual curated names")
            return create_multilingual_curated_names(selected_languages, exclude_common_words)
    
    else:
        print("   [ERROR] FACEBOOK DATABASE UNAVAILABLE: names-dataset package not installed")
        print("   [FALLBACK] Using multilingual curated names")
        return create_multilingual_curated_names(selected_languages, exclude_common_words)
        print("   ❌ FACEBOOK DATABASE UNAVAILABLE: names-dataset package not installed")
        print("   🔄 Using curated basic names")
        return create_basic_english_names()
        print("   🔄 Using enhanced basic English names (~150 names)")
        return create_basic_english_names()

def download_name_database():
    """Legacy function - now redirects to package-based loading."""
    print("   Note: Using names-dataset package instead of direct download")
    return None

def filter_problematic_names(names):
    """Remove names that coincide with common English words."""
    
    # Comprehensive list of common English words that should not be masked
    common_words = {
        # Colors
        'red', 'blue', 'green', 'yellow', 'orange', 'purple', 'pink', 'brown', 'black', 'white', 'grey', 'gray',
        'violet', 'rose', 'lily', 'amber', 'jade', 'ruby', 'pearl', 'ivory', 'silver', 'golden',
        
        # Months/Time
        'january', 'february', 'march', 'april', 'may', 'june', 'july', 'august', 'september', 'october', 'november', 'december',
        'jan', 'feb', 'mar', 'apr', 'jun', 'jul', 'aug', 'sep', 'oct', 'nov', 'dec',
        'monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday',
        'morning', 'evening', 'dawn', 'dusk',
        
        # Common verbs/actions
        'will', 'can', 'may', 'might', 'shall', 'should', 'would', 'could',
        'go', 'come', 'run', 'walk', 'talk', 'see', 'hear', 'feel', 'think',
        'hope', 'wish', 'want', 'need', 'have', 'get', 'make', 'take', 'give',
        'work', 'play', 'read', 'write', 'draw', 'paint', 'sing', 'dance',
        
        # Common nouns
        'house', 'home', 'car', 'book', 'phone', 'computer', 'table', 'chair', 'door', 'window',
        'water', 'food', 'money', 'time', 'day', 'night', 'week', 'month', 'year',
        'man', 'woman', 'child', 'person', 'people', 'family', 'friend',
        'school', 'work', 'job', 'office', 'shop', 'store', 'market',
        'hill', 'valley', 'river', 'lake', 'sea', 'ocean', 'mountain', 'forest', 'field',
        'stone', 'rock', 'wood', 'tree', 'flower', 'grass', 'leaf',
        
        # Common adjectives
        'good', 'bad', 'big', 'small', 'long', 'short', 'tall', 'wide', 'narrow',
        'hot', 'cold', 'warm', 'cool', 'dry', 'wet', 'clean', 'dirty',
        'new', 'old', 'young', 'fast', 'slow', 'hard', 'soft', 'strong', 'weak',
        'happy', 'sad', 'angry', 'calm', 'quiet', 'loud', 'bright', 'dark',
        'rich', 'poor', 'free', 'busy', 'easy', 'hard', 'simple', 'difficult',
        
        # Occupations/Titles
        'doctor', 'nurse', 'teacher', 'student', 'worker', 'manager', 'director',
        'cook', 'baker', 'farmer', 'driver', 'pilot', 'captain', 'judge',
        'artist', 'writer', 'singer', 'actor', 'player', 'hunter', 'fisher',
        'mason', 'taylor', 'turner', 'walker', 'parker', 'porter', 'carter',
        
        # Geographic/Place terms
        'north', 'south', 'east', 'west', 'center', 'middle', 'top', 'bottom',
        'city', 'town', 'village', 'country', 'state', 'place', 'area', 'region',
        'street', 'road', 'avenue', 'lane', 'way', 'path', 'bridge', 'park',
        'church', 'temple', 'hall', 'tower', 'castle', 'palace',
        'island', 'beach', 'shore', 'coast', 'port', 'bay', 'gulf',
        
        # Animals
        'cat', 'dog', 'bird', 'fish', 'horse', 'cow', 'pig', 'sheep', 'goat',
        'lion', 'tiger', 'bear', 'wolf', 'fox', 'deer', 'rabbit', 'mouse',
        'bee', 'ant', 'fly', 'spider', 'snake', 'frog', 'duck', 'swan',
        
        # Body parts
        'head', 'face', 'eye', 'ear', 'nose', 'mouth', 'hand', 'foot', 'arm', 'leg',
        'heart', 'brain', 'bone', 'skin', 'hair', 'nail',
        
        # Common words that are often names
        'angel', 'joy', 'grace', 'faith', 'hope', 'charity', 'patience', 'mercy',
        'peace', 'love', 'dream', 'wish', 'star', 'moon', 'sun', 'sky',
        'rain', 'snow', 'wind', 'storm', 'thunder', 'lightning',
        
        # Misc common words
        'name', 'names', 'word', 'letter', 'number', 'page', 'line', 'text',
        'picture', 'image', 'photo', 'video', 'music', 'song', 'sound', 'voice',
        'game', 'sport', 'team', 'player', 'winner', 'loser', 'score', 'point',
        'question', 'answer', 'problem', 'solution', 'idea', 'plan', 'goal',
        'start', 'end', 'begin', 'finish', 'stop', 'pause', 'continue',
        'yes', 'no', 'maybe', 'never', 'always', 'sometimes', 'often', 'rarely',
        
        # Single letters that might be names
        'a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm',
        'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z',
        
        # Common conversation words
        'the', 'okay', 'yeah', 'well', 'but', 'move', 'thank', 'thanks', 'please',
        'sorry', 'excuse', 'hello', 'hi', 'bye', 'goodbye', 'welcome', 'congrats',
        'um', 'uh', 'oh', 'ah', 'hmm', 'wow', 'great', 'nice', 'fine', 'sure',
        'actually', 'really', 'basically', 'obviously', 'definitely', 'probably',
        'maybe', 'perhaps', 'anyway', 'however', 'therefore', 'because', 'since',
        'although', 'though', 'unless', 'until', 'while', 'during', 'before', 'after',
        
        # Additional problematic words
        'whale', 'whales', 'look', 'let', 'lets', 'every', 'put', 'nor', 'did', 'wave',
        'wait', 'mummy', 'baby', 'lots', 'lovely', 'say', 'painting', 'shower', 'nowhere',
        'pardon', 'breakfast', 'bible', 'wayo', 'shovel', 'shadows', 'meow', 'noy',
        
        # Question words and pronouns
        'what', 'who', 'when', 'where', 'why', 'how', 'which', 'whose',
        'he', 'she', 'it', 'they', 'we', 'you', 'i', 'me', 'him', 'her', 'them', 'us',
        'this', 'that', 'these', 'those', 'here', 'there', 'now', 'then',
        'some', 'any', 'all', 'each', 'every', 'many', 'much', 'few', 'little',
        
        # Prepositions and conjunctions 
        'in', 'on', 'at', 'by', 'for', 'with', 'without', 'to', 'from', 'of',
        'about', 'above', 'below', 'under', 'over', 'through', 'between', 'among',
        'and', 'or', 'but', 'so', 'if', 'as', 'than', 'like', 'unlike',
        
        # Common verbs that might appear as names
        'do', 'does', 'did', 'be', 'am', 'is', 'are', 'was', 'were', 'been', 'being',
        'have', 'has', 'had', 'having', 'get', 'got', 'getting', 'put', 'putting',
        'say', 'said', 'saying', 'tell', 'told', 'telling', 'ask', 'asked', 'asking',
        'know', 'knew', 'known', 'knowing', 'think', 'thought', 'thinking',
        'look', 'looked', 'looking', 'seem', 'seemed', 'seeming', 'try', 'tried', 'trying',
        
        # Common contractions (without apostrophes since they're stripped in processing)
        'hes', 'shes', 'its', 'thats', 'whats', 'wheres', 'theres', 'youre', 'theyre',
        'were', 'werent', 'cant', 'dont', 'wont', 'isnt', 'arent', 'wasnt', 'havent', 'hasnt', 'hadnt',
        'ill', 'youll', 'hell', 'shell', 'well', 'theyll', 'ive', 'youve', 'weve', 'theyve',
        'id', 'youd', 'hed', 'shed', 'wed', 'theyd', 'lets', 'theres', 'heres',
        
        # Numbers and quantities
        'one', 'two', 'three', 'four', 'five', 'six', 'seven', 'eight', 'nine', 'ten',
        'first', 'second', 'third', 'fourth', 'fifth', 'last', 'next', 'another',
        'more', 'most', 'less', 'least', 'enough', 'too', 'very', 'quite', 'rather',
        
        # Technology and modern words
        'app', 'web', 'site', 'email', 'text', 'call', 'chat', 'post', 'like', 'share',
        'click', 'tap', 'swipe', 'scroll', 'zoom', 'search', 'find', 'save', 'delete',
        'update', 'download', 'upload', 'install', 'connect', 'wifi', 'internet',
        
        # Business and finance
        'buy', 'sell', 'pay', 'cost', 'price', 'cheap', 'expensive', 'free', 'sale',
        'deal', 'offer', 'discount', 'tax', 'fee', 'bill', 'check', 'cash', 'card',
        'bank', 'account', 'loan', 'debt', 'invest', 'profit', 'loss', 'budget'
    }
    
    # Remove problematic names
    filtered = names - common_words
    
    # Also remove very short names (1-2 characters) as they're likely to be false positives
    filtered = {name for name in filtered if len(name) >= 3}
    
    return filtered

def create_multilingual_curated_names(selected_languages=None, exclude_common_words=True):
    """Create comprehensive multilingual curated names database from CSV file.
    
    Loads names from data/curated_names.csv which contains columns:
    - name: The personal name
    - type: Either 'first' or 'last' (surname)
    - language: Language code (english, chinese, french, german, hindi, spanish, italian, arabic, polynesian)
    
    Args:
        selected_languages (list): List of language codes to include. 
                                 If None, includes all languages.
                                 Supported: ['english', 'chinese', 'french', 'german', 'hindi', 'spanish', 'italian', 'arabic', 'polynesian']
    
    Returns:
        tuple: (first_names_set, surnames_set) with combined names from selected languages
    """
    import csv
    import os
    
    # Default to all languages if none specified
    if selected_languages is None:
        selected_languages = ['english', 'chinese', 'french', 'german', 'hindi', 'spanish', 'italian', 'arabic', 'polynesian']
    
    print(f"   [DATABASE] Loading multilingual curated names from CSV")
    print(f"   [LANGUAGES] Selected: {', '.join(selected_languages)}")
    
    # Locate the CSV file (relative to this script)
    script_dir = os.path.dirname(os.path.abspath(__file__))
    csv_path = os.path.join(script_dir, 'data', 'curated_names.csv')
    
    if not os.path.exists(csv_path):
        print(f"   [ERROR] CSV file not found at: {csv_path}")
        print(f"   [FALLBACK] Using empty name sets")
        return set(), set()
    
    # Load names from CSV and organize by language and type
    language_stats = {lang: {'first': 0, 'last': 0} for lang in selected_languages}
    all_first_names = set()
    all_surnames = set()
    
    try:
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                name = row['name'].strip().lower()
                name_type = row['type'].strip().lower()
                language = row['language'].strip().lower()
                
                # Only include names from selected languages
                if language not in selected_languages:
                    continue
                
                # Add to appropriate set
                if name_type == 'first':
                    all_first_names.add(name)
                    language_stats[language]['first'] += 1
                elif name_type == 'last':
                    all_surnames.add(name)
                    language_stats[language]['last'] += 1
        
        # Report statistics by language
        for lang in selected_languages:
            stats = language_stats[lang]
            print(f"   [LOADED] {lang.title()}: {stats['first']} first names, {stats['last']} surnames")
        
        print(f"   [SUCCESS] CSV database loaded: {len(all_first_names)} first names, {len(all_surnames)} surnames")
        
        # Conditionally filter out common English words to prevent false positives
        if exclude_common_words:
            print(f"   [FILTERING] Removing common English words from curated database...")
            filtered_first_names = filter_problematic_names(all_first_names)
            filtered_surnames = filter_problematic_names(all_surnames)
            
            removed_first = len(all_first_names) - len(filtered_first_names)
            removed_last = len(all_surnames) - len(filtered_surnames)
            
            if removed_first > 0 or removed_last > 0:
                print(f"   [REMOVED] {removed_first} common first names, {removed_last} common surnames")
            
            print(f"   [FINAL] {len(filtered_first_names)} first names, {len(filtered_surnames)} surnames after filtering")
            
            return filtered_first_names, filtered_surnames
        else:
            print(f"   [INFO] Common word filtering disabled - using full curated database")
            return all_first_names, all_surnames
        
    except Exception as e:
        print(f"   [ERROR] Failed to load CSV file: {e}")
        print(f"   [FALLBACK] Using empty name sets")
        return set(), set()


def create_basic_english_names():
    """Legacy function - now redirects to multilingual database with English only."""
    print("   [LEGACY] create_basic_english_names() called - redirecting to multilingual database")
    return create_multilingual_curated_names(['english'], exclude_common_words=True)  # Always filter for English

def create_enhanced_name_masker(use_facebook_names=False, use_facebook_surnames=False, selected_languages=None, exclude_common_words=True, excluded_names=None):
    """Create enhanced name masking with global database and logging.
    
    Args:
        use_facebook_names (bool): If True, use Facebook's global first names database.
                                 If False, use curated multilingual first names only.
        use_facebook_surnames (bool): If True, use Facebook's global surnames database.
                                    If False, use curated multilingual surnames only.
        selected_languages (list): Languages to include in curated database.
        exclude_common_words (bool): If True, filter out common English words from name database.
                                    Default: True (recommended for English transcripts).
        excluded_names (set): Set of names (case-insensitive) to exclude from masking.
                            Useful for preserving public figures, celebrities, or specific participants.
    """
    
    print("   Loading global name database...")
    names_data = load_global_names(use_facebook_names, use_facebook_surnames, selected_languages, exclude_common_words)
    
    # Handle both old single-set return and new tuple return
    if isinstance(names_data, tuple):
        first_names, surnames = names_data
    else:
        # Fallback for old format
        first_names = names_data
        surnames = set()
    
    # Apply excluded names (case-insensitive)
    if excluded_names:
        excluded_names_lower = {name.lower() for name in excluded_names}
        first_names = {name for name in first_names if name.lower() not in excluded_names_lower}
        surnames = {name for name in surnames if name.lower() not in excluded_names_lower}
        print(f"   ℹ️  Excluding {len(excluded_names)} name(s) from masking")
    
    # Title prefixes
    titles = {'mr', 'mrs', 'ms', 'miss', 'dr', 'prof', 'professor', 'sir', 'madam', 'lord', 'lady'}
    
    # Initialize replacement log
    replacement_log = []
    
    def mask_names_with_logging(text, filename="unknown"):
        """Enhanced name masking function with comprehensive logging."""
        nonlocal replacement_log
        replacement_log = []  # Reset for new file
        
        sentences = re.split(r'(?<=[.!?])\s+', text)
        masked_sentences = []
        replacement_counter = 1
        
        for sentence_idx, sentence in enumerate(sentences):
            if not sentence.strip():
                masked_sentences.append(sentence)
                continue
                
            words = sentence.split()
            masked_words = []
            
            for word_idx, word in enumerate(words):
                original_word = word
                clean_word = re.sub(r'[^\w]', '', word.lower())
                
                # Skip if empty after cleaning
                if not clean_word:
                    masked_words.append(word)
                    continue
                
                # Check if word is capitalized
                is_capitalized = word[0].isupper() if word else False
                
                # Check if it's at the start of a sentence (excluding titles/pronouns)
                is_sentence_start = word_idx == 0
                
                # Check if it's a name (requires capitalization, but be careful at sentence start)
                is_first_name = clean_word in first_names
                is_surname = clean_word in surnames
                is_title = clean_word in titles
                
                # Only mask if:
                # 1. It's capitalized AND in the database, AND
                # 2. Either it's NOT at sentence start, OR it appears multiple times capitalized
                should_mask = is_capitalized and (is_first_name or is_surname)
                
                # Special handling: if at sentence start, only mask if it's uncommon or appears mid-sentence too
                # This prevents masking "Will you..." but allows masking repeated names
                if should_mask and is_sentence_start:
                    # Keep common words at sentence start (they're likely not names)
                    # The filter already removed most common words, but double-check context
                    should_mask = True  # Trust the filtered database
                
                # Check for title + name combinations
                if is_title and word_idx + 1 < len(words):
                    next_word = re.sub(r'[^\w]', '', words[word_idx + 1].lower())
                    next_is_capitalized = words[word_idx + 1][0].isupper() if words[word_idx + 1] else False
                    if next_is_capitalized and (next_word in first_names or next_word in surnames):
                        # Log the title
                        replacement_log.append({
                            'order': replacement_counter,
                            'original': original_word,
                            'replacement': '[TITLE]',
                            'context_sentence': sentence.strip(),
                            'filename': filename
                        })
                        replacement_counter += 1
                        masked_words.append('[TITLE]')
                        continue
                
                # Process names (requires capitalization + database match)
                if should_mask:
                    # Preserve punctuation
                    punctuation = re.findall(r'[^\w]', word)
                    
                    # Choose appropriate replacement based on name type
                    if is_surname:
                        replacement = '[SURNAME]' + ''.join(punctuation)
                    else:
                        replacement = '[NAME]' + ''.join(punctuation)
                    
                    # Log the replacement
                    replacement_log.append({
                        'order': replacement_counter,
                        'original': original_word,
                        'replacement': replacement,
                        'context_sentence': sentence.strip(),
                        'filename': filename
                    })
                    replacement_counter += 1
                    
                    masked_words.append(replacement)
                else:
                    masked_words.append(word)
            
            masked_sentences.append(' '.join(masked_words))
        
        return ' '.join(masked_sentences)
    
    def get_replacement_log():
        """Get the current replacement log."""
        return replacement_log

    # Return both the masking function and the log function
    return mask_names_with_logging, get_replacement_log

def save_replacement_log(replacement_log, filename):
    """Save replacement log to CSV file in root directory."""
    if not replacement_log:
        return
    
    # Create logs directory in root (same level as transcription.py)
    logs_dir = Path("name_masking_logs")
    logs_dir.mkdir(parents=True, exist_ok=True)
    
    # Create log filename
    base_name = Path(filename).stem
    log_file = logs_dir / f"name_replacements_{base_name}.csv"
    
    # Write CSV
    with open(log_file, 'w', newline='', encoding='utf-8') as csvfile:
        fieldnames = ['filename', 'order', 'original', 'replacement', 'context_sentence']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        
        writer.writeheader()
        for entry in replacement_log:
            writer.writerow(entry)
    
    print(f"   📝 Name replacement log saved: {log_file}")
    print(f"   📊 Total replacements: {len(replacement_log)}")

# Legacy function for backwards compatibility
def create_core_name_masker(use_facebook_database=False):
    """Legacy function - redirects to enhanced version."""
    # Convert old boolean to new separate parameters
    # Use default exclude_common_words=True for backwards compatibility
    mask_func, _ = create_enhanced_name_masker(use_facebook_database, use_facebook_database, exclude_common_words=True)
    return mask_func


def remove_emojis_and_unicode_artifacts(text):
    """
    Remove emojis, Unicode artifacts, and non-English characters from transcription.
    Ensures clean, readable text without visual artifacts.
    """
    import re
    
    # Step 1: Remove emoji patterns (comprehensive Unicode ranges)
    # Emoticons and misc symbols
    emoji_pattern = re.compile(
        "["
        u"\U0001F600-\U0001F64F"  # emoticons
        u"\U0001F300-\U0001F5FF"  # symbols & pictographs
        u"\U0001F680-\U0001F6FF"  # transport & map symbols
        u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
        u"\U00002702-\U000027B0"  # dingbats
        u"\U000024C2-\U0001F251"  # enclosed characters
        u"\U0001F900-\U0001F9FF"  # supplemental symbols
        u"\U0001FA70-\U0001FAFF"  # symbols and pictographs extended-A
        u"\U00002600-\U000026FF"  # miscellaneous symbols
        u"\U00002700-\U000027BF"  # dingbats
        "]+", 
        flags=re.UNICODE
    )
    text = emoji_pattern.sub(r'', text)
    
    # Step 2: Remove other problematic Unicode characters and symbols
    # Remove various Unicode symbols that shouldn't be in speech transcription
    text = re.sub(r'[^\x00-\x7F\s]', '', text)  # Remove non-ASCII characters
    
    # Step 3: Remove isolated symbols and artifacts
    text = re.sub(r'\s+[^\w\s\[\].,!?;:\'"()-]+\s+', ' ', text)  # Remove isolated symbols
    
    # Step 4: Clean up multiple spaces
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()

def detect_and_fix_repetitions(text, max_repetitions=5):
    """
    Enhanced repetition detection and fixing for AI-generated artifacts.
    Handles massive repetitive patterns that are clearly AI artifacts.
    This addresses patterns like 'I have this. I have this.' repeated 80+ times.
    """
    import re

    # Step 1: Detect and fix sentence-level repetitions (AI artifacts)
    # Pattern: Same sentence repeated many times
    sentences = re.split(r'(?<=[.!?])\s+', text)
    cleaned_sentences = []
    
    i = 0
    while i < len(sentences):
        current_sentence = sentences[i].strip()
        if not current_sentence:
            i += 1
            continue
            
        # Count consecutive identical sentences
        repetition_count = 1
        j = i + 1
        while j < len(sentences) and sentences[j].strip() == current_sentence:
            repetition_count += 1
            j += 1
        
        # If we find massive repetitions (more than max_repetitions), it's likely an AI artifact
        if repetition_count > max_repetitions:
            # Keep only one instance and add a note about the repetition
            cleaned_sentences.append(current_sentence)
            if repetition_count > 10:  # Only note for truly massive repetitions
                cleaned_sentences.append("[...]")
            i = j
        else:
            # Keep all instances of normal repetitions
            for k in range(repetition_count):
                cleaned_sentences.append(current_sentence)
            i = j
    
    text = ' '.join(cleaned_sentences)
    
    # Step 2: Detect and fix phrase-level repetitions within sentences
    # Pattern: "word word word word" repeated many times within a sentence
    def fix_phrase_repetitions(match):
        full_match = match.group(0)
        repeated_phrase = match.group(1)
        
        # Count how many times the phrase is repeated
        parts = full_match.split(repeated_phrase)
        repetition_count = len([p for p in parts if not p.strip()]) + full_match.count(repeated_phrase)
        
        if repetition_count > max_repetitions * 2:  # More lenient for phrases
            return f"{repeated_phrase} [...]"
        else:
            return full_match
    
    # Look for patterns like "phrase phrase phrase phrase"
    # This regex finds a phrase (2-6 words) repeated multiple times
    phrase_pattern = r'(\b\w+(?:\s+\w+){1,5})\s*(?:\1\s*){4,}'
    text = re.sub(phrase_pattern, fix_phrase_repetitions, text, flags=re.IGNORECASE)
    
    # Step 3: Remove excessive counting sequences that are causing the bug
    counting_pattern = r'(?:one\s+two\s+three\s+four\s+five\s+six\s+seven\s+eight\s+nine\s*){3,}'
    text = re.sub(counting_pattern, '[...]', text, flags=re.IGNORECASE)

    # Step 4: Remove excessive emoji repetitions
    text = re.sub(r'([^\w\s])\s*\1{5,}', r'\1\1\1', text)

    # Step 5: Clean up simple word repetitions more conservatively
    words = text.split()
    cleaned_words = []
    i = 0
    while i < len(words):
        word = words[i]
        count = 1

        # Count consecutive identical words
        while i + count < len(words) and words[i + count].lower() == word.lower():
            count += 1

        # Keep only reasonable number of repetitions
        if count > max_repetitions:
            cleaned_words.extend(words[i:i+2])  # Keep just 2 instances
        else:
            cleaned_words.extend(words[i:i+count])

        i += count

    result = ' '.join(cleaned_words)
    return result if result else text

def clean_transcription_text(text):
    """
    Comprehensive text cleaning for transcription output.
    Addresses all 8 formatting issues from the original function.
    """
    
    # 1. Handle undefined terms and speech artifacts
    text = re.sub(r'\s+undefined\s+', ' [undefined] ', text, flags=re.IGNORECASE)
    text = re.sub(r'\s+undefine\s+', ' [undefined] ', text, flags=re.IGNORECASE)
    text = re.sub(r'\s+undefin\s+', ' [undefined] ', text, flags=re.IGNORECASE)
    text = re.sub(r'\s+undefining\s+', ' [undefined] ', text, flags=re.IGNORECASE)
    
    # 2. Handle filler words and hesitations (mark but don't remove)
    text = re.sub(r'\b(um|uh|er|ah|hmm|mm|mhm|uhuh|uhhuh)\b', r'[\1]', text, flags=re.IGNORECASE)
    
    # 3. Fix common transcription errors and contractions
    contractions = {
        # Basic contractions
        r'\bim\b': "I'm",
        r'\byoure\b': "you're", 
        r'\btheyre\b': "they're",
        r'\bwere\b': "we're",
        r'\bitsits\b': "it's",
        r'\bits\b': "it's",  # context-dependent, be careful
        r'\bdont\b': "don't",
        r'\bwont\b': "won't",
        r'\bcant\b': "can't",
        r'\bisnt\b': "isn't",
        r'\barent\b': "aren't",
        r'\bwasnt\b': "wasn't",
        r'\bwerent\b': "weren't",
        r'\bhasnt\b': "hasn't",
        r'\bhavent\b': "haven't",
        r'\bhadnt\b': "hadn't",
        r'\bcouldnt\b': "couldn't",
        r'\bshouldnt\b': "shouldn't",
        r'\bwouldnt\b': "wouldn't",
        r'\bdidnt\b': "didn't",
        r'\bdoesnt\b': "doesn't",

        # Space-separated contractions (Whisper model artifacts)
        r'\bisn\s+t\b': "isn't",
        r'\bdoesn\s+t\b': "doesn't",
        r'\bdon\s+t\b': "don't",
        r'\bwon\s+t\b': "won't",
        r'\bcan\s+t\b': "can't",
        r'\baren\s+t\b': "aren't",
        r'\bwasn\s+t\b': "wasn't",
        r'\bweren\s+t\b': "weren't",
        r'\bhasn\s+t\b': "hasn't",
        r'\bhaven\s+t\b': "haven't",
        r'\bhadn\s+t\b': "hadn't",
        r'\bcouldn\s+t\b': "couldn't",
        r'\bshouldn\s+t\b': "shouldn't",
        r'\bwouldn\s+t\b': "wouldn't",
        r'\bdidn\s+t\b': "didn't",
        
        # Will contractions  
        r'\btheyll\b': "they'll",
        r'\byoull\b': "you'll",
        r'\bitll\b': "it'll",
        r'\bthatll\b': "that'll",
        r'\bwholl\b': "who'll",
        r'\bwhatll\b': "what'll",
        
        # Have contractions
        r'\bive\b': "I've",
        r'\byouve\b': "you've",
        r'\bweve\b': "we've",
        r'\btheyve\b': "they've",
        r'\bcouldve\b': "could've",
        r'\bshouldve\b': "should've",
        r'\bwouldve\b': "would've",
        
        # Would contractions
        r'\bid\b': "I'd",
        r'\byoud\b': "you'd", 
        r'\bhed\b': "he'd",
        r'\bshed\b': "she'd",
        r'\bwed\b': "we'd",
        r'\btheyd\b': "they'd",
    }
    
    # Apply contraction fixes (case-insensitive)
    for pattern, replacement in contractions.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    
    # 4. Fix common misheard words and phrases
    common_fixes = {
        r'\band all\b': 'and all',
        r'\bkind of\b': 'kind of',
        r'\bsort of\b': 'sort of', 
        r'\byou know\b': 'you know',
        r'\bi mean\b': 'I mean',
        r'\bi think\b': 'I think',
        r'\bi guess\b': 'I guess',
        r'\bi feel like\b': 'I feel like',
        r'\bgoing to\b': 'going to',
        r'\bwant to\b': 'want to',
        r'\bhave to\b': 'have to',
        r'\btrying to\b': 'trying to',
        r'\bused to\b': 'used to',
        r'\bsupposed to\b': 'supposed to',
        r'\ba lot of\b': 'a lot of',
        r'\ba little bit\b': 'a little bit',
        r'\ball of\b': 'all of',
        r'\bsome of\b': 'some of',
        r'\bone of\b': 'one of',
        r'\bnone of\b': 'none of',
    }
    
    for pattern, replacement in common_fixes.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    
    # 5. Fix spacing and punctuation issues
    text = re.sub(r'\s+', ' ', text)  # Multiple spaces to single
    text = re.sub(r'\s+([,.])', r'\1', text)  # Space before punctuation
    text = re.sub(r'([,.])([a-zA-Z])', r'\1 \2', text)  # Missing space after punctuation
    
    # 6. Capitalize sentence beginnings and "I"
    # Split into sentences and capitalize first word
    sentences = re.split(r'([.!?])', text)
    for i in range(0, len(sentences), 2):  # Every other element is sentence content
        if sentences[i].strip():
            # Capitalize first word
            sentences[i] = re.sub(r'^\s*([a-z])', lambda m: m.group(1).upper(), sentences[i])
            # Capitalize standalone "i"
            sentences[i] = re.sub(r'\bi\b', 'I', sentences[i])
    
    text = ''.join(sentences)
    # 8. Fix punctuation spacing issues
    # Fix spaces before punctuation (keep space after)
    text = re.sub(r'\s+([.,!?;:])', r'\1', text)
    
    # Ensure space after punctuation marks (but not if followed by another punctuation)
    text = re.sub(r'([.,!?;:])([^\s.,!?;:])', r'\1 \2', text)
    
    # Fix multiple punctuation issues
    text = re.sub(r'([.,!?])\s*\1+', r'\1', text)  # Remove duplicate punctuation
    text = re.sub(r'([.,!?;:])\s+([.,!?;:])', r'\1\2', text)  # Fix space between punctuation
    
    # Fix quotation mark spacing
    text = re.sub(r'\s+"', '"', text)  # Remove space before opening quote
    text = re.sub(r'"\s+', '" ', text)  # Ensure space after closing quote
    text = re.sub(r'"\s*([.,!?;:])', r'"\1', text)  # No space between quote and punctuation
    
    # Fix parentheses spacing  
    text = re.sub(r'\s*\(\s*', ' (', text)  # Space before opening paren
    text = re.sub(r'\s*\)\s*', ') ', text)  # Space after closing paren
    text = re.sub(r'\(\s+', '(', text)  # No space after opening paren
    text = re.sub(r'\s+\)', ')', text)  # No space before closing paren

    
    # 7. Clean up extra whitespace and return
    text = text.strip()
    return text

def format_text_with_line_breaks(text, line_length=130):
    """Format text without line wrapping (updated to keep continuous flow)"""
    return text.strip()


def create_word_document(text_content, output_path, input_file, mask_names=False, model_name=None, language=None, fix_repetitions=False, original_filename=None, use_facebook_names=False, use_facebook_surnames=False, enhance_audio_enabled=False, speaker_attribution=False):
    """Create a Word document from text content with proper title and metadata."""
    if not DOCX_AVAILABLE:
        print("   ⚠️  Word document creation skipped - python-docx not available")
        return
        
    try:
        doc = Document()
        
        # Extract filename without path and extension for title
        # Use original_filename if provided (to avoid temp filenames)
        import os
        if original_filename:
            filename = os.path.splitext(os.path.basename(original_filename))[0]
        else:
            filename = os.path.splitext(os.path.basename(input_file))[0]

        # Add title with filename (size 15 equivalent to heading level 1)
        title = doc.add_heading(filename, 0)
        title.alignment = 1  # Center alignment

        # Add metadata section (same as TXT file but without the removed Features line)
        metadata_para = doc.add_paragraph()
        metadata_para.add_run("="*52).bold = True
        
        # Show original filename with extension (not temporary filename)
        if original_filename:
            input_filename = os.path.basename(original_filename)
        else:
            input_filename = os.path.basename(input_file)
        doc.add_paragraph(f"Input file: {input_filename}")
        doc.add_paragraph(f"Processed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        if model_name:
            doc.add_paragraph(f"Model: {model_name}")
        
        # Individual parameter lines in plain English (only when enabled)
        # Order: Audio enhancement -> Language -> Repetitions -> Privacy -> Speaker attribution
        if enhance_audio_enabled:
            doc.add_paragraph(f"Audio enhanced before transcription")
        if language:
            doc.add_paragraph(f"Language: {language.title()}")
        if fix_repetitions:
            doc.add_paragraph(f"Edited to remove likely spurious repetitions.")
        if mask_names:
            # Determine which masking system was used
            if use_facebook_names or use_facebook_surnames:
                doc.add_paragraph(f"Privacy: Personal names masked using Facebook list (review recommended)")
            else:
                doc.add_paragraph(f"Privacy: Personal names masked using internal list (review recommended)")
        if speaker_attribution:
            doc.add_paragraph(f"Speaker attribution: Enabled (accuracy depends on recording quality, speaker count, and speaker similarity)")
        
        metadata_para2 = doc.add_paragraph()
        metadata_para2.add_run("="*52).bold = True

        # Add blank line before transcription
        doc.add_paragraph()

        # Add the main content
        # Split by double newlines to preserve paragraph structure
        paragraphs = text_content.split("\n\n")
        for paragraph in paragraphs:
            if paragraph.strip():
                para = doc.add_paragraph(paragraph.strip())
                # Set line spacing to 1.5
                para.paragraph_format.line_spacing = 1.5

        # Save the document
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return False
def enhance_audio(input_path, output_path):
    """Enhance audio quality for better transcription."""
    print(f" Enhancing audio quality...")

    # Load audio
    audio, sr = librosa.load(input_path, sr=16000)

    # Apply interview_clarity preset enhancements
    # 1. Amplification (3.61x gain)
    audio = audio * 3.61

    # 2. Noise reduction using spectral subtraction
    stft = librosa.stft(audio)
    magnitude = np.abs(stft)
    phase = np.angle(stft)

    # Estimate noise from first 0.5 seconds
    noise_frames = int(0.5 * sr / 512)
    noise_magnitude = np.mean(magnitude[:, :noise_frames], axis=1, keepdims=True)

    # Spectral subtraction
    alpha = 2.0
    clean_magnitude = magnitude - alpha * noise_magnitude
    clean_magnitude = np.maximum(clean_magnitude, 0.1 * magnitude)

    # Reconstruct audio
    clean_stft = clean_magnitude * np.exp(1j * phase)
    audio = librosa.istft(clean_stft)

    # 4. Dynamic range compression
    threshold = 0.1
    ratio = 4.0
    audio_abs = np.abs(audio)
    mask = audio_abs > threshold
    audio[mask] = np.sign(audio[mask]) * (threshold + (audio_abs[mask] - threshold) / ratio)

    # 5. Normalization to prevent clipping
    max_val = np.max(np.abs(audio))
    if max_val > 0.95:
        audio = audio * (0.95 / max_val)

    # Save enhanced audio
    sf.write(output_path, audio, sr)
    print(f"   Enhanced audio saved to {output_path}")

def perform_speaker_diarization(audio_path):
    """
    Perform speaker diarization using pyannote.audio.
    Returns a list of (start_time, end_time, speaker_label) tuples.
    Requires HuggingFace token with pyannote model access.
    """
    print(f"  🎭 Performing speaker diarization...")
    
    # Try to import pyannote - if it fails, skip diarization entirely
    try:
        from pyannote.audio import Pipeline
        import torch
    except (ImportError, OSError) as e:
        print(f"   ⚠️  Pyannote import failed: {str(e)}")
        print(f"   Skipping speaker diarization...")
        return None
    
    try:
        
        # Check for HuggingFace token
        hf_token = os.environ.get("HF_TOKEN") or os.environ.get("HUGGINGFACE_TOKEN")
        if not hf_token:
            # Try to read from hf_token.txt file
            token_file = "hf_token.txt"
            if os.path.exists(token_file):
                with open(token_file, 'r') as f:
                    hf_token = f.read().strip()
        
        if not hf_token:
            print(f"   ⚠️  Warning: No HuggingFace token found!")
            print(f"   Set HF_TOKEN environment variable or create hf_token.txt file")
            print(f"   Skipping speaker diarization...")
            return None
        
        # Load diarization pipeline
        print(f"   Loading pyannote/speaker-diarization-3.1 model...")
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        
        # Try new API first (token), fallback to old API (use_auth_token)
        try:
            pipeline = Pipeline.from_pretrained(
                "pyannote/speaker-diarization-3.1",
                use_auth_token=hf_token
            )
        except TypeError:
            # Newer versions use 'token' instead of 'use_auth_token'
            pipeline = Pipeline.from_pretrained(
                "pyannote/speaker-diarization-3.1",
                token=hf_token
            )
        
        pipeline.to(device)
        print(f"   Model loaded on {device}")
        
        # Perform diarization
        print(f"   Running diarization (this may take a while)...")
        diarization = pipeline(audio_path)
        
        # Extract speaker segments
        segments = []
        for turn, _, speaker in diarization.itertracks(yield_label=True):
            segments.append({
                'start': turn.start,
                'end': turn.end,
                'speaker': speaker
            })
        
        print(f"   ✅ Diarization complete: {len(segments)} speaker segments identified")
        print(f"   Detected speakers: {len(set(seg['speaker'] for seg in segments))}")
        
        return segments
        
    except ImportError as e:
        print(f"   ⚠️  pyannote.audio not installed: {e}")
        print(f"   Install with: pip install pyannote.audio")
        print(f"   Skipping speaker diarization...")
        return None
    except Exception as e:
        print(f"   ⚠️  Diarization failed: {e}")
        print(f"   Continuing without speaker attribution...")
        return None

def save_transcript_file(output_file, formatted_text, base_name, input_file, 
                        original_filename, model, language, mask_names, fix_repetitions,
                        use_facebook_names, use_facebook_surnames, enhance_audio_enabled, 
                        speaker_attribution=False):
    """
    Save transcript to file with appropriate headers.
    
    Args:
        output_file: Path to output file
        formatted_text: Formatted transcript text
        base_name: Base name for display
        input_file: Original input file path
        original_filename: Original filename (if different from input_file)
        model: Model used for transcription
        language: Language specified (if any)
        mask_names: Whether name masking was enabled
        fix_repetitions: Whether repetition fixing was enabled
        use_facebook_names: Whether Facebook names database was used
        use_facebook_surnames: Whether Facebook surnames database was used
        enhance_audio_enabled: Whether audio enhancement was enabled
        speaker_attribution: Whether this file has speaker labels
    """
    with open(output_file, 'w', encoding='utf-8') as f:
        # Add title with filename
        display_filename = Path(original_filename).stem if original_filename else base_name
        f.write(f"{display_filename}\n")
        f.write(f"{'='*len(display_filename)}\n\n")
        
        # Show original filename with extension
        if original_filename:
            input_filename = Path(original_filename).name
        else:
            input_filename = Path(input_file).name
        f.write(f"Input file: {input_filename}\n")
        f.write(f"Processed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        if model:
            f.write(f"Model: {model}\n")
        
        # Individual parameter lines
        # Order: Audio enhancement -> Language -> Repetitions -> Privacy -> Speaker attribution
        if enhance_audio_enabled:
            f.write(f"Audio enhanced before transcription\n")
        if language:
            f.write(f"Language: {language.title()}\n")
        if fix_repetitions:
            f.write(f"Edited to remove likely spurious repetitions.\n")
        if mask_names:
            if use_facebook_names or use_facebook_surnames:
                f.write(f"Privacy: Personal names masked using Facebook list (review recommended)\n")
            else:
                f.write(f"Privacy: Personal names masked using internal list (review recommended)\n")
        if speaker_attribution:
            f.write(f"Speaker attribution: Enabled (accuracy depends on recording quality, speaker count, and speaker similarity)\n")
        
        f.write(f"{'='*len(display_filename)}\n\n")
        
        # Write transcription
        f.write(formatted_text + "\n\n")

def apply_speaker_labels_to_transcript(transcript_text, diarization_segments, audio_duration):
    """
    Apply speaker labels to transcript text based on diarization segments.
    Uses a simple approach of breaking transcript into roughly equal time segments.
    
    Note: This is an approximation since we don't have word-level timestamps.
    For more accurate attribution, would need word-level timestamps from Whisper.
    """
    if not diarization_segments:
        return transcript_text
    
    print(f"   Applying speaker labels to transcript...")
    
    # Split transcript into words
    words = transcript_text.split()
    if not words:
        return transcript_text
    
    # Estimate time per word (rough approximation)
    time_per_word = audio_duration / len(words) if len(words) > 0 else 0
    
    # Build attributed transcript
    result_lines = []
    current_speaker = None
    current_text = []
    word_index = 0
    
    for word in words:
        # Estimate current timestamp
        current_time = word_index * time_per_word
        
        # Find which speaker segment this corresponds to
        speaker_for_word = None
        for seg in diarization_segments:
            if seg['start'] <= current_time <= seg['end']:
                speaker_for_word = seg['speaker']
                break
        
        # If speaker changed, start new line
        if speaker_for_word and speaker_for_word != current_speaker:
            # Save previous speaker's text
            if current_text:
                result_lines.append(f"[{current_speaker}] {' '.join(current_text)}")
            current_speaker = speaker_for_word
            current_text = [word]
        else:
            current_text.append(word)
        
        word_index += 1
    
    # Add final speaker's text
    if current_text and current_speaker:
        result_lines.append(f"[{current_speaker}] {' '.join(current_text)}")
    
    attributed_transcript = '\n\n'.join(result_lines)
    print(f"   ✅ Speaker labels applied")
    
    return attributed_transcript

def transcribe_audio_only(audio_path, model="openai/whisper-large-v3", language=None):
    """Perform transcription only - no diarization."""
    print(f"  Transcribing audio (no diarization)...")

    # Initialize transcription pipeline
    try:
        device = 0 if torch.cuda.is_available() else -1
        
        # Base pipeline arguments
        pipeline_args = {
            "model": model,  # User-specified or default model
            "dtype": torch.float16 if torch.cuda.is_available() else torch.float32,
            "device": device,
            "return_timestamps": False,    # No timestamps needed
            "chunk_length_s": 15,         # Process in 15-second chunks
            "stride_length_s": 2,
            "ignore_warning": True        # 2-second overlap between chunks
        }
        
        # Add language specification if requested
        if language:
            pipeline_args["generate_kwargs"] = {"language": language}
            print(f"   Language set to {language.title()}")
        
        transcriber = pipeline("automatic-speech-recognition", **pipeline_args)
        print(f"   Transcription model loaded (device: {'GPU' if device >= 0 else 'CPU'})")
    except Exception as e:
        print(f"   Failed to load transcription model: {e}")
        return None

    # Perform transcription
    print("    Running transcription...")
    try:
        # Load audio with librosa instead of letting transformers use ffmpeg
        audio_array, sr = librosa.load(audio_path, sr=16000)
        transcription_result = transcriber(audio_array)
        print(f"     Transcription completed")
        return transcription_result
    except Exception as e:
        print(f"   Transcription failed: {e}")
        return None

def process_audio_file(input_file, mask_names=False, model="openai/whisper-large-v3", language=None, original_filename=None, fix_repetitions=False, output_name=None, save_name_masking_logs=False, save_enhanced_audios=False, enhance_audio_enabled=False, use_facebook_names=False, use_facebook_surnames=False, selected_languages=None, speaker_attribution=False, exclude_common_words=True, excluded_names=None):
    """Process a single audio file with ultimate enhanced transcription."""

    print(f"\n{'='*80}")
    print(f" PROCESSING: {input_file}")
    print(f"{'='*80}")

    # Check if audio processing libraries are available
    if not AUDIO_LIBS_AVAILABLE or not TRANSFORMERS_AVAILABLE:
        print(f" ❌ Error: Required audio processing libraries not available")
        print(f"    Audio libraries: {'✅' if AUDIO_LIBS_AVAILABLE else '❌'}")
        print(f"    Transformers: {'✅' if TRANSFORMERS_AVAILABLE else '❌'}")
        print(f"    Please install missing dependencies or use the proper environment")
        return False

    if not os.path.exists(input_file):
        print(f" Error: Audio file not found: {input_file}")
        return False

    try:
        # Step 1: Audio enhancement
        print("\n🎵 AUDIO ENHANCEMENT")
        print("-" * 40)

        # Create enhanced audio filename using original filename if available
        if original_filename:
            base_name = Path(original_filename).stem
        else:
            input_path = Path(input_file)
            base_name = input_path.stem
        
        # Conditionally enhance audio based on enhance_audio_enabled flag
        audio_file_for_transcription = input_file  # Default to original file
        enhanced_audio_path = None
        
        if enhance_audio_enabled:
            if save_enhanced_audios:
                os.makedirs("output/enhanced_audio", exist_ok=True)
                enhanced_audio_path = f"output/enhanced_audio/{base_name}_enhanced.wav"
                print(f"   🎧 Saving enhanced audio to: {enhanced_audio_path}")
            else:
                # Use temporary file if not saving enhanced audio
                import tempfile
                temp_fd, enhanced_audio_path = tempfile.mkstemp(suffix='.wav', prefix='enhanced_')
                os.close(temp_fd)  # Close file descriptor, keep path
                print(f"   🎧 Using temporary enhanced audio (not saved)")

            enhance_audio(input_file, enhanced_audio_path)
            audio_file_for_transcription = enhanced_audio_path
        else:
            print(f"   🎧 Using original audio (enhancement disabled)")

        # Step 2: Speaker Diarization (if enabled)
        diarization_segments = None
        audio_duration = None
        if speaker_attribution:
            print("\n🎭 SPEAKER DIARIZATION")
            print("-" * 40)
            diarization_segments = perform_speaker_diarization(audio_file_for_transcription)
            
            # Get audio duration for speaker label mapping
            try:
                import librosa
                audio_duration = librosa.get_duration(path=audio_file_for_transcription)
                print(f"   Audio duration: {audio_duration:.2f} seconds")
            except Exception as e:
                print(f"   ⚠️  Could not determine audio duration: {e}")

        # Step 3: Transcription 
        print("\n📝 TRANSCRIPTION")
        print("-" * 40)
        transcription_result = transcribe_audio_only(audio_file_for_transcription, model, language)

        if not transcription_result:
            print(f"❌ Transcription failed - cannot continue")
            if enhance_audio_enabled and save_enhanced_audios:
                print(f"   🎧 Enhanced audio preserved at: {enhanced_audio_path}")
            elif enhance_audio_enabled and enhanced_audio_path:
                # Clean up temporary file
                try:
                    os.unlink(enhanced_audio_path)
                except:
                    pass
            return False

        # Step 4: Process and clean transcription text
        print("\n🔧 TEXT PROCESSING AND ENHANCEMENT")
        print("-" * 40)
        
        # Get the main transcription text
        main_text = transcription_result.get('text', '')
        
        # Store base transcription for dual output when speaker attribution is enabled
        base_transcript_text = main_text
        
        # Apply speaker attribution if enabled and segments available
        if speaker_attribution and diarization_segments and audio_duration:
            print("\n  🎭 SPEAKER ATTRIBUTION")
            print("-" * 40)
            main_text = apply_speaker_labels_to_transcript(main_text, diarization_segments, audio_duration)
        elif speaker_attribution:
            print("\n  ⚠️  Speaker attribution requested but diarization unavailable")
        
        # Initialize name masking if requested
        mask_names_func = None
        get_replacement_log_func = None
        if mask_names:
            print("   Initializing enhanced name masking with global database...")
            mask_names_func, get_replacement_log_func = create_enhanced_name_masker(use_facebook_names, use_facebook_surnames, selected_languages, exclude_common_words, excluded_names)

        # Remove emojis and Unicode artifacts first
        print("   Removing emojis and Unicode artifacts...")
        main_text = remove_emojis_and_unicode_artifacts(main_text)
        # Apply all text processing steps BEFORE line wrapping
        
        if fix_repetitions:
            print("   Detecting and fixing repetitive patterns...")
            main_text = detect_and_fix_repetitions(main_text)
        else:
            print("   Skipping repetition detection (not enabled)")
        
        print("   Applying comprehensive spelling corrections...")
        main_text = clean_transcription_text(main_text)
        
        if mask_names_func:
            print("   Applying enhanced name masking with logging...")
            # Use the enhanced masking function with filename for logging
            filename = Path(original_filename).name if original_filename else Path(input_file).name
            main_text = mask_names_func(main_text, filename)
            
            # Save the replacement log
            if get_replacement_log_func:
                replacement_log = get_replacement_log_func()
                if replacement_log and save_name_masking_logs:
                    print(f"   💾 Saving name replacement log ({len(replacement_log)} replacements)")
                    save_replacement_log(replacement_log, filename)
                elif replacement_log and not save_name_masking_logs:
                    print(f"   ℹ️  Name replacement log available ({len(replacement_log)} replacements) - use --save-name-masking-logs to save")
        
        # Now apply line wrapping to the cleaned text
        print("   Applying line wrapping...")
        formatted_main_text = format_text_with_line_breaks(main_text)

        # Step 4: Save results with enhanced formatting
        print("\n💾 SAVING RESULTS")
        print("-" * 40)

        # Create output filename using custom name, original filename, or input filename
        if output_name:
            # Use custom output name
            base_name = output_name
            print(f"📝 Using custom output name: {output_name}")
        elif original_filename:
            # Use original filename (for temp file processing)
            base_name = Path(original_filename).stem
        else:
            # Use input filename
            input_path = Path(input_file)
            base_name = input_path.stem
        
        # Ensure output directory exists
        os.makedirs("output/transcripts", exist_ok=True)
        
        # If speaker attribution is enabled, save both versions
        if speaker_attribution and diarization_segments and audio_duration:
            # Save version WITHOUT speaker labels
            output_file_base = f"output/transcripts/{base_name}_transcript.txt"
            
            # Process base transcript (without speakers) through same cleaning pipeline
            print("   Processing base transcript (without speaker labels)...")
            base_text_processed = base_transcript_text
            
            # Apply same text processing to base transcript
            base_text_processed = remove_emojis_and_unicode_artifacts(base_text_processed)
            if fix_repetitions:
                base_text_processed = detect_and_fix_repetitions(base_text_processed)
            base_text_processed = clean_transcription_text(base_text_processed)
            if mask_names_func:
                filename = Path(original_filename).name if original_filename else Path(input_file).name
                base_text_processed = mask_names_func(base_text_processed, filename)
            formatted_base_text = format_text_with_line_breaks(base_text_processed)
            
            # Save base version
            save_transcript_file(output_file_base, formatted_base_text, base_name, input_file, 
                               original_filename, model, language, mask_names, fix_repetitions,
                               use_facebook_names, use_facebook_surnames, enhance_audio_enabled, 
                               speaker_attribution=False)  # Mark as no speakers for header
            print(f"   ✅ Base transcript saved: {output_file_base}")
            
            # Save version WITH speaker labels  
            output_file_speakers = f"output/transcripts/{base_name}_transcript_with_speakers.txt"
            save_transcript_file(output_file_speakers, formatted_main_text, base_name, input_file,
                               original_filename, model, language, mask_names, fix_repetitions,
                               use_facebook_names, use_facebook_surnames, enhance_audio_enabled,
                               speaker_attribution=True)  # Mark as with speakers for header
            print(f"   ✅ Speaker-attributed transcript saved: {output_file_speakers}")
            
            # Create Word documents for both versions
            word_output_base = output_file_base.replace('.txt', '.docx')
            create_word_document(formatted_base_text, word_output_base, input_file, mask_names, model, 
                               language, fix_repetitions, original_filename, use_facebook_names, 
                               use_facebook_surnames, enhance_audio_enabled, speaker_attribution=False)
            print(f"   ✅ Base Word document saved: {word_output_base}")
            
            word_output_speakers = output_file_speakers.replace('.txt', '.docx')
            create_word_document(formatted_main_text, word_output_speakers, input_file, mask_names, model,
                               language, fix_repetitions, original_filename, use_facebook_names,
                               use_facebook_surnames, enhance_audio_enabled, speaker_attribution=True)
            print(f"   ✅ Speaker-attributed Word document saved: {word_output_speakers}")
            
        else:
            # Standard single output (no speaker attribution or it failed)
            output_file = f"output/transcripts/{base_name}_transcript.txt"
            
            # Save single transcript
            save_transcript_file(output_file, formatted_main_text, base_name, input_file,
                               original_filename, model, language, mask_names, fix_repetitions,
                               use_facebook_names, use_facebook_surnames, enhance_audio_enabled,
                               speaker_attribution)
            print(f"   ✅ Transcript saved: {output_file}")
            
            # Create Word document
            word_output_file = output_file.replace('.txt', '.docx')
            create_word_document(formatted_main_text, word_output_file, input_file, mask_names, model, 
                               language, fix_repetitions, original_filename, use_facebook_names, 
                               use_facebook_surnames, enhance_audio_enabled, speaker_attribution)
            print(f"   ✅ Word document saved: {word_output_file}")

        # Cleanup
        if enhance_audio_enabled and save_enhanced_audios:
            print(f"   🎧 Enhanced audio preserved at: {enhanced_audio_path}")
        elif enhance_audio_enabled and enhanced_audio_path:
            # Clean up temporary enhanced audio file
            try:
                os.unlink(enhanced_audio_path)
                print(f"   🗑️  Temporary enhanced audio cleaned up")
            except:
                pass

        print(f"\n✅ SUCCESS: Enhanced processing completed for {input_file}")
        return True

    except Exception as e:
        print(f" Error processing {input_file}: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Main function to process a single audio file with ultimate enhancements."""
    parser = argparse.ArgumentParser(description='Ultimate Enhanced Transcription Processor')
    parser.add_argument('audio_file', help='Input audio file path')
    parser.add_argument('--mask-personal-names', action='store_true',
                       help='Enable personal name masking (boolean flag)')
    parser.add_argument('--use-facebook-names-for-masking', action='store_true',
                       help='Use Facebook global first names database (730K+ names, may cause false positives). Default: use curated basic names only.')
    parser.add_argument('--use-facebook-surnames-for-masking', action='store_true',
                       help='Use Facebook global surnames database (980K+ surnames, may cause false positives). Default: use curated basic surnames only.')
    parser.add_argument('--languages-for-name-masking', type=str, nargs='*', 
                       default=['english', 'chinese', 'french', 'german', 'hindi', 'spanish'],
                       help='Languages to include in curated names database. Supported: english, chinese, french, german, hindi, spanish. Default: all languages.')
    parser.add_argument('--exclude-common-english-words-from-name-masking', action='store_true', default=None,
                       help='Exclude common English words (e.g., "will", "long", "art") from name masking. Default: ON when --language english, otherwise OFF.')
    parser.add_argument('--exclude-names-from-masking', type=str, default=None,
                       help='Comma-separated list of names to exclude from masking (e.g., "John,Sarah,Einstein"). Case-insensitive.')
    parser.add_argument('--exclude-names-file', type=str, default=None,
                       help='Path to file containing names to exclude from masking (one name per line). Case-insensitive.')
    parser.add_argument('--model', type=str, default="openai/whisper-large-v3",
                       help='HuggingFace model ID (default: openai/whisper-large-v3)')
    parser.add_argument('--language', type=str, default=None,
                       help='Specify transcription language (e.g., "english", "spanish", "french"). Default: auto-detect language')
    parser.add_argument('--original-filename', type=str, default=None,
                       help='Original filename to use for output (for temp file processing)')
    parser.add_argument('--fix-spurious-repetitions', action='store_true',
                       help='Enable automatic repetition pattern detection and removal')
    parser.add_argument('--output-name', type=str, default=None,
                       help='Custom name for output files (without extension)')
    parser.add_argument('--single-file', type=str, default=None,
                       help=argparse.SUPPRESS)  # Hidden argument for internal HPC use only
    parser.add_argument('--save-name-masking-logs', action='store_true',
                       help='Save detailed logs of name replacements to name_masking_logs/ directory')
    parser.add_argument('--save-enhanced-audio', action='store_true',
                       help='Save enhanced audio files to output/enhanced_audio/ directory')
    parser.add_argument('--enhance-audio', action='store_true',
                       help='Enable audio enhancement for better transcription quality (disabled by default)')
    parser.add_argument('--speaker-attribution', action='store_true',
                       help='Enable speaker diarization/attribution (disabled by default). Note: Accuracy depends on recording quality, number of speakers, and speaker similarity.')
    
    args = parser.parse_args()
    
    mask_names = args.mask_personal_names
    use_facebook_names = args.use_facebook_names_for_masking
    use_facebook_surnames = args.use_facebook_surnames_for_masking
    selected_languages = args.languages_for_name_masking
    force_english = args.language == "english" if args.language else False
    transcription_language = args.language
    original_filename = args.original_filename
    fix_repetitions = args.fix_spurious_repetitions
    output_name = args.output_name
    single_file_path = args.single_file
    save_name_masking_logs = args.save_name_masking_logs
    save_enhanced_audios = args.save_enhanced_audio
    enhance_audio_enabled = args.enhance_audio
    speaker_attribution_enabled = args.speaker_attribution
    
    # Determine if common English words should be excluded from name masking
    # Default: ON when --language english, otherwise OFF
    if args.exclude_common_english_words_from_name_masking is None:
        # User didn't specify, use smart default
        exclude_common_words = (args.language == "english")
    else:
        # User explicitly specified
        exclude_common_words = args.exclude_common_english_words_from_name_masking
    
    # Parse excluded names from command line and file
    excluded_names = set()
    if args.exclude_names_from_masking:
        # Parse comma-separated names (case-insensitive)
        excluded_names.update(name.strip().lower() for name in args.exclude_names_from_masking.split(',') if name.strip())
    
    if args.exclude_names_file:
        # Read names from file (one per line, case-insensitive)
        try:
            with open(args.exclude_names_file, 'r', encoding='utf-8') as f:
                excluded_names.update(line.strip().lower() for line in f if line.strip())
        except FileNotFoundError:
            print(f"⚠️  Warning: Exclusion file not found: {args.exclude_names_file}")
        except Exception as e:
            print(f"⚠️  Warning: Could not read exclusion file: {e}")
    
    print(" ENHANCED TRANSCRIPTION PROCESSOR")
    print("=" * 55)
    
    # Determine the input file to process
    input_file = single_file_path if single_file_path else args.audio_file
    
    print(f"Input: {input_file}")
    print(f"Model: {args.model}")
    if transcription_language:
        print(f"Language: {transcription_language.title()}")
    else:
        print("Language: Auto-detect")
    features = "No line wrapping, no timestamps"
    if fix_repetitions:
        features += ", repetition detection"
    print(f"Features: {features}")
    print(f"Text Processing: Comprehensive spelling corrections")
    if mask_names:
        print("Privacy: Personal name masking enabled")
        if excluded_names:
            print(f"         Excluding {len(excluded_names)} name(s) from masking")
    if speaker_attribution_enabled:
        print("Speaker Attribution: Enabled (requires HuggingFace token)")

    success = process_audio_file(input_file, mask_names, args.model, transcription_language, original_filename, fix_repetitions, output_name, save_name_masking_logs, save_enhanced_audios, enhance_audio_enabled, use_facebook_names, use_facebook_surnames, selected_languages, speaker_attribution_enabled, exclude_common_words, excluded_names)

    if success:
        print("\n JOB COMPLETED SUCCESSFULLY!")
        sys.exit(0)
    else:
        print("\n JOB FAILED!")
        sys.exit(1)

if __name__ == "__main__":
    main()
